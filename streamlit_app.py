import streamlit as st
from io import BytesIO
from docx import Document
import inference_openai as ioai
import os

st.set_page_config(page_title="AI Document Editor", layout="wide")
st.title("AI-Powered Document Editor")
st.write("Upload any .docx; only the text in Normal-style paragraphs is AI-rewritten in place. The rest of the template stays exactly as you provided it.")

uploaded = st.file_uploader("Choose a .docx file", type="docx")
if uploaded:
    # Load the client's exact template
    doc = Document(uploaded)

    # Iterate over every paragraph in the document, in order
    for para in doc.paragraphs:
        # Only target body text—style must be "Normal"
        if para.style.name == "Normal" and para.text.strip():
            original = para.text
            try:
                edited = ioai.generate_edit(original)
            except Exception as e:
                st.warning(f"AI edit failed on a paragraph—using original. Error: {e}")
                edited = original

            # Remove all existing runs, then write back the edited text
            for run in para.runs:
                para._element.remove(run._element)
            para.add_run(edited)
            para.style = doc.styles["Normal"]  # re-apply style

    # Save back to memory
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button(
        "Download Edited Document",
        data=buf,
        file_name="edited_with_template.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
