#!/usr/bin/env python3
import os
import sys
import openai
from docx import Document
from docx.shared import Pt

# 1️⃣ Configure
openai.api_key = os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    print("❌ ERROR: Set the OPENAI_API_KEY environment variable first.")
    sys.exit(1)

MODEL = "gpt-3.5-turbo"  # or your preferred deployment
TEMPERATURE = 0.7
TOP_P = 0.9

def rewrite_paragraph(text: str) -> str:
    """Call OpenAI to rewrite a single paragraph."""
    if not text.strip():
        return text  # leave empty or whitespace-only paragraphs unchanged

    prompt = [
        {"role": "system", "content": (
            "You are an expert grant editor. "
            "Rewrite the following paragraph for clarity, conciseness, and impact, "
            "without changing its meaning."
        )},
        {"role": "user", "content": text}
    ]
    resp = openai.ChatCompletion.create(
        model=MODEL,
        messages=prompt,
        temperature=TEMPERATURE,
        top_p=TOP_P,
    )
    return resp.choices[0].message.content.strip()

def process_doc(input_path: str, output_path: str):
    """Read input_path, rewrite each paragraph, save to output_path."""
    doc = Document(input_path)
    new_doc = Document()
    
    # Copy document-level styles (optional; python-docx does much by default)
    new_doc.styles = doc.styles

    for para in doc.paragraphs:
        # Preserve the entire paragraph style (heading level, bullets, indentation)
        new_para = new_doc.add_paragraph()
        new_para.style = para.style
        
        # Get the raw text
        original = para.text
        # Get rewritten text
        edited = rewrite_paragraph(original)
        
        # Rebuild runs to preserve bold/italic/etc if needed:
        # For simplicity, we strip formatting and put the full edited text as one run:
        run = new_para.add_run(edited)
        run.font.size = para.runs[0].font.size if para.runs and para.runs[0].font.size else Pt(11)
        # (If you need to preserve bold/italic within runs, you'd need a more complex merge.)

    new_doc.save(output_path)
    print(f"✅ Saved edited document to {output_path}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python edit_docx.py <input.docx>")
        sys.exit(1)
    inp = sys.argv[1]
    base, ext = os.path.splitext(inp)
    outp = f"{base}.edited{ext}"
    process_doc(inp, outp)
